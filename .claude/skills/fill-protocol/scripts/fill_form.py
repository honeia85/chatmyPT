#!/usr/bin/env python3
"""Fill a Korean Word form template while preserving styles, tables, fonts, and page layout.

Core principles (DO NOT BREAK):
1. Always open existing template via Document(path) — never create from scratch.
2. Modify cell/paragraph TEXT only. Preserve all run-level styles.
3. Apply cantSplit to every row that gets filled (prevents page-break-mid-row).
4. Set Korean font with eastAsia attribute (run.font.name alone fails for Korean).
5. Validate: report empty cells and paragraphs that didn't match.
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.table import _Cell


DEFAULT_KOREAN_FONT = "맑은 고딕"


# ---------- Style preservation helpers ----------

def _set_run_korean_font(run, font_name: str) -> None:
    """Set font for a run including eastAsia attribute (mandatory for Hangul)."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)


def _apply_cant_split(row) -> None:
    """Mark row to never split across pages."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _make_blank_paragraph() -> "OxmlElement":
    """Create an empty paragraph that renders as a single Enter press.

    Forces single line height (line=240) and zero spacing-before/after,
    so the blank line is exactly one body-text line tall — never inflates
    the document's apparent line spacing.
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)
    p.append(pPr)
    return p


def _replace_paragraph_text_keep_style(para: Paragraph, new_text: str,
                                        korean_font: str | None = None) -> None:
    """Replace the entire text content of a paragraph while keeping its style.

    Strategy: keep the first run's properties as the template style. Remove all
    other runs. Replace the first run's text with new_text. For multi-line
    content, split on \n and use w:br between lines (within same run-style block).
    """
    # Capture template run (first one) style by copying its rPr
    runs = para.runs
    template_rPr = None
    if runs:
        template_run_elem = runs[0]._element
        rPr = template_run_elem.find(qn("w:rPr"))
        if rPr is not None:
            template_rPr = rPr

    # Remove all existing runs
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)

    # Add new run with the captured style
    new_run = OxmlElement("w:r")
    if template_rPr is not None:
        # Deep copy template rPr
        from copy import deepcopy
        new_run.append(deepcopy(template_rPr))

    # Split on \n — insert w:br between lines, w:t for text segments
    lines = new_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement("w:br")
            new_run.append(br)
        if line:
            t = OxmlElement("w:t")
            t.text = line
            t.set(qn("xml:space"), "preserve")
            new_run.append(t)

    para._element.append(new_run)

    if korean_font:
        # Reapply Korean font to the new run
        from docx.text.run import Run
        run_obj = Run(new_run, para)
        _set_run_korean_font(run_obj, korean_font)


def _replace_cell_text(cell: _Cell, new_text: str,
                        korean_font: str | None = None) -> None:
    """Replace a cell's text content. Use the first paragraph as template."""
    if not cell.paragraphs:
        # Cell has no paragraph — add one
        cell.add_paragraph(new_text)
        if korean_font:
            for r in cell.paragraphs[0].runs:
                _set_run_korean_font(r, korean_font)
        return

    # Replace first paragraph, then remove the rest
    template_para = cell.paragraphs[0]

    # If the new content has multiple lines, we replace first paragraph
    # with the first line, and add additional paragraphs for remaining lines.
    lines = new_text.split("\n")

    _replace_paragraph_text_keep_style(template_para, lines[0],
                                        korean_font=korean_font)

    # Remove all paragraphs after the first
    for p in list(cell._tc.findall(qn("w:p")))[1:]:
        cell._tc.remove(p)

    # Add new paragraphs for remaining lines (cloning first paragraph's pPr)
    if len(lines) > 1:
        from copy import deepcopy
        first_p = cell._tc.find(qn("w:p"))
        first_pPr = first_p.find(qn("w:pPr")) if first_p is not None else None
        first_rPr = None
        first_r = first_p.find(qn("w:r")) if first_p is not None else None
        if first_r is not None:
            first_rPr = first_r.find(qn("w:rPr"))

        for line in lines[1:]:
            new_p = OxmlElement("w:p")
            if first_pPr is not None:
                new_p.append(deepcopy(first_pPr))
            new_r = OxmlElement("w:r")
            if first_rPr is not None:
                new_r.append(deepcopy(first_rPr))
            t = OxmlElement("w:t")
            t.text = line
            t.set(qn("xml:space"), "preserve")
            new_r.append(t)
            new_p.append(new_r)
            cell._tc.append(new_p)

        if korean_font:
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_korean_font(r, korean_font)


# ---------- FormFiller class ----------

@dataclass
class FillResult:
    matched: list[str] = field(default_factory=list)
    unmatched: list[str] = field(default_factory=list)


class FormFiller:
    def __init__(self, template_path: str | Path,
                 korean_font: str = DEFAULT_KOREAN_FONT,
                 blank_between_paragraphs: bool = True,
                 blank_around_section_header: bool = True,
                 blank_around_all_section_headers: bool = False,
                 normalize_page_breaks: bool = True):
        self.path = Path(template_path).expanduser().resolve()
        if not self.path.exists():
            raise FileNotFoundError(self.path)
        self.doc = Document(str(self.path))
        self.korean_font = korean_font
        self.blank_between_paragraphs = blank_between_paragraphs
        self.blank_around_section_header = blank_around_section_header
        self.blank_around_all_section_headers = blank_around_all_section_headers
        self.normalize_page_breaks_flag = normalize_page_breaks
        self._filled_rows: set[int] = set()
        self._table_results = FillResult()
        self._paragraph_results = FillResult()

    # ---- Table cell filling ----

    def _cell_text(self, cell: _Cell) -> str:
        return "\n".join(p.text for p in cell.paragraphs).strip()

    def _label_match(self, cell_text: str, label: str) -> bool:
        # Normalize whitespace and newlines
        norm_cell = re.sub(r"\s+", "", cell_text)
        norm_label = re.sub(r"\s+", "", label)
        return norm_cell == norm_label

    def fill_table_kv(self, label: str, value: str) -> bool:
        """Find a cell whose text == label, fill the next cell on the right.

        Returns True if filled, False otherwise.
        Skips merged duplicate cells (same _tc reference).
        """
        for table in self.doc.tables:
            for row_idx, row in enumerate(table.rows):
                # Track unique cells in this row (skip merged duplicates)
                seen_tcs: set[int] = set()
                cells_in_row: list[_Cell] = []
                for c in row.cells:
                    if id(c._tc) not in seen_tcs:
                        seen_tcs.add(id(c._tc))
                        cells_in_row.append(c)

                for ci, cell in enumerate(cells_in_row):
                    if self._label_match(self._cell_text(cell), label):
                        # Found label cell. Fill the next cell on the right.
                        if ci + 1 < len(cells_in_row):
                            target = cells_in_row[ci + 1]
                            _replace_cell_text(target, value,
                                               korean_font=self.korean_font)
                            _apply_cant_split(row)
                            self._table_results.matched.append(label)
                            return True
        self._table_results.unmatched.append(label)
        return False

    # ---- Paragraph (section) filling ----

    def replace_paragraphs_after(self, header_text: str, new_content: str,
                                  stop_pattern: str | None = None) -> bool:
        """Find a paragraph matching header_text, then replace all paragraphs
        between this header and the next section header (or stop_pattern) with
        new_content.

        new_content is split by \n\n into separate paragraphs (preserving the
        style of the first replaced paragraph).
        """
        body = self.doc.element.body
        all_ps = list(self.doc.paragraphs)

        # Find header paragraph
        header_idx = None
        for i, p in enumerate(all_ps):
            if self._label_match(p.text, header_text):
                header_idx = i
                break

        if header_idx is None:
            self._paragraph_results.unmatched.append(header_text)
            return False

        # Determine end paragraph (next numbered section header or stop_pattern)
        if stop_pattern:
            end_re = re.compile(stop_pattern)
        else:
            # Match patterns like "1. ", "2. ", ... "18. "
            end_re = re.compile(r"^\s*\d+\.\s+\S")

        end_idx = len(all_ps)
        for i in range(header_idx + 1, len(all_ps)):
            if end_re.match(all_ps[i].text):
                end_idx = i
                break

        # Paragraphs to replace: header_idx+1 .. end_idx-1
        # Strategy: replace first paragraph in range, remove rest, add new paragraphs
        if header_idx + 1 >= end_idx:
            # No paragraphs between header and next section — just insert
            from copy import deepcopy
            template_p = all_ps[header_idx]._element
            template_pPr = template_p.find(qn("w:pPr"))
            template_r = template_p.find(qn("w:r"))
            template_rPr = template_r.find(qn("w:rPr")) if template_r is not None else None

            insert_after = template_p
            # Blank line right after section header
            if self.blank_around_section_header:
                blank_p = _make_blank_paragraph()
                insert_after.addnext(blank_p)
                insert_after = blank_p
            chunks = new_content.split("\n\n")
            for ci, chunk in enumerate(chunks):
                if ci > 0 and self.blank_between_paragraphs:
                    blank_p = _make_blank_paragraph()
                    insert_after.addnext(blank_p)
                    insert_after = blank_p
                new_p = OxmlElement("w:p")
                # New paragraph should NOT have header style — use default (no pPr)
                new_r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = chunk
                t.set(qn("xml:space"), "preserve")
                new_r.append(t)
                new_p.append(new_r)
                insert_after.addnext(new_p)
                # Apply Korean font
                from docx.text.run import Run
                _set_run_korean_font(Run(new_r, None), self.korean_font)
                insert_after = new_p
            # Blank line right before next section header
            if self.blank_around_section_header:
                blank_p = _make_blank_paragraph()
                insert_after.addnext(blank_p)
            self._paragraph_results.matched.append(header_text)
            return True

        # Replace first paragraph in range
        first_target = all_ps[header_idx + 1]
        chunks = new_content.split("\n\n")
        _replace_paragraph_text_keep_style(first_target, chunks[0],
                                           korean_font=self.korean_font)

        # Remove all paragraphs after first_target up to end_idx
        for i in range(header_idx + 2, end_idx):
            p_elem = all_ps[i]._element
            p_elem.getparent().remove(p_elem)

        # Insert blank paragraph right after section header (before first body)
        first_target_elem = first_target._element
        if self.blank_around_section_header:
            blank_p = _make_blank_paragraph()
            first_target_elem.addprevious(blank_p)

        # Add additional chunks as new paragraphs after first_target
        from copy import deepcopy
        first_pPr = first_target_elem.find(qn("w:pPr"))
        first_r = first_target_elem.find(qn("w:r"))
        first_rPr = first_r.find(qn("w:rPr")) if first_r is not None else None

        insert_after = first_target_elem
        for chunk in chunks[1:]:
            if self.blank_between_paragraphs:
                blank_p = _make_blank_paragraph()
                insert_after.addnext(blank_p)
                insert_after = blank_p
            new_p = OxmlElement("w:p")
            if first_pPr is not None:
                new_p.append(deepcopy(first_pPr))
            new_r = OxmlElement("w:r")
            if first_rPr is not None:
                new_r.append(deepcopy(first_rPr))
            t = OxmlElement("w:t")
            t.text = chunk
            t.set(qn("xml:space"), "preserve")
            new_r.append(t)
            new_p.append(new_r)
            insert_after.addnext(new_p)
            from docx.text.run import Run
            _set_run_korean_font(Run(new_r, None), self.korean_font)
            insert_after = new_p

        # Blank line right before next section header
        if self.blank_around_section_header:
            blank_p = _make_blank_paragraph()
            insert_after.addnext(blank_p)

        self._paragraph_results.matched.append(header_text)
        return True

    # ---- Single-paragraph in-place text replace ----

    def replace_paragraph_matching(self, matcher: str, new_text: str,
                                    mode: str = "startswith") -> bool:
        """Replace the entire text of the first paragraph that matches.

        mode: 'startswith' | 'contains' | 'exact'
        Preserves the paragraph's pPr and the first run's rPr (style).
        """
        for p in self.doc.paragraphs:
            text = p.text
            ok = False
            if mode == "startswith":
                ok = text.startswith(matcher)
            elif mode == "contains":
                ok = matcher in text
            elif mode == "exact":
                ok = text.strip() == matcher.strip()
            if ok:
                _replace_paragraph_text_keep_style(p, new_text,
                                                    korean_font=self.korean_font)
                self._paragraph_results.matched.append(f"<para>{matcher}")
                return True
        self._paragraph_results.unmatched.append(f"<para>{matcher}")
        return False

    # ---- Document-wide passes ----

    def apply_blank_around_all_section_headers(self) -> int:
        """Scan all top-level paragraphs and add blank lines above and below
        every numbered section header (e.g. '1. ', '12. ').

        OPT-IN ONLY. Use this when the institutional review will tolerate
        layout drift (page count change). For strict form-fidelity submissions,
        leave disabled (default) and rely on per-section blanks added during
        replace_paragraphs_after().

        Skips:
        - Headers whose previous sibling is already an empty paragraph
          (avoids double-blanks when section was filled via section_replace)
        - Headers whose next sibling is already an empty paragraph
        - Paragraphs inside tables (only top-level body paragraphs scanned)

        Returns the number of blank paragraphs inserted.
        """
        header_re = re.compile(r"^\s*\d+\.\s+\S")
        body = self.doc.element.body
        # Collect all top-level <w:p> elements (skip those inside <w:tbl>)
        all_top_ps = [el for el in body if el.tag == qn("w:p")]
        inserted = 0

        def is_blank(p_elem) -> bool:
            if p_elem is None or p_elem.tag != qn("w:p"):
                return False
            # Empty if no <w:t> with text content
            for t in p_elem.iter(qn("w:t")):
                if t.text and t.text.strip():
                    return False
            return True

        def text_of(p_elem) -> str:
            return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))

        for p_elem in all_top_ps:
            text = text_of(p_elem)
            if not header_re.match(text):
                continue
            prev = p_elem.getprevious()
            nxt = p_elem.getnext()
            if not is_blank(prev):
                p_elem.addprevious(_make_blank_paragraph())
                inserted += 1
            if not is_blank(nxt):
                p_elem.addnext(_make_blank_paragraph())
                inserted += 1
        return inserted

    # ---- Validation & save ----

    def validate(self) -> list[str]:
        warnings: list[str] = []
        for label in self._table_results.unmatched:
            warnings.append(f"[TABLE-MISS] Label not found: {label!r}")
        for header in self._paragraph_results.unmatched:
            warnings.append(f"[SECTION-MISS] Header not found: {header!r}")
        return warnings

    def report(self) -> str:
        n_table_ok = len(self._table_results.matched)
        n_table_miss = len(self._table_results.unmatched)
        n_para_ok = len(self._paragraph_results.matched)
        n_para_miss = len(self._paragraph_results.unmatched)
        return (
            f"Filled {n_table_ok} table cells, {n_para_ok} sections.\n"
            f"Missed: {n_table_miss} cells, {n_para_miss} sections."
        )

    def normalize_page_breaks(self) -> int:
        """Remove dangling empty paragraphs whose sole content is a page break,
        and transfer the break to the next content paragraph via pageBreakBefore.

        Why: templates often place `<w:p><w:r><w:br w:type="page"/></w:r></w:p>`
        after a table or section header to force the next block onto a new page.
        When the preceding content's height varies (e.g. an abstract table grows
        with content), the empty paragraph can spill onto a page by itself and
        the page break then forces the next block one more page forward —
        producing a visibly blank page.

        Replacing this pattern with `<w:pageBreakBefore/>` on the next content
        paragraph's `pPr` preserves the "start on a new page" intent regardless
        of where the preceding content ends, eliminating the blank page.

        Returns the number of paragraphs normalized.
        """
        from copy import deepcopy  # noqa: F401 (kept for parity with other helpers)

        body = self.doc.element.body
        children = list(body)
        fixed = 0

        for i, el in enumerate(children):
            if not el.tag.endswith("}p"):
                continue
            # Only collapse paragraphs with NO real text, containing a page break
            text = "".join((t.text or "") for t in el.iter(qn("w:t")))
            if text.strip():
                continue
            page_brs = [b for b in el.iter(qn("w:br"))
                        if b.get(qn("w:type")) == "page"]
            if not page_brs:
                continue
            # Find the next sibling content paragraph (non-empty p or table)
            target = None
            for j in range(i + 1, len(children)):
                sib = children[j]
                if sib.tag.endswith("}p"):
                    sib_text = "".join((t.text or "") for t in sib.iter(qn("w:t")))
                    if sib_text.strip():
                        target = sib
                        break
                elif sib.tag.endswith("}tbl"):
                    # A table has no pPr; leave the break alone.
                    target = None
                    break
            if target is None:
                continue
            # Attach pageBreakBefore to target's pPr (idempotent)
            pPr = target.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                target.insert(0, pPr)
            if pPr.find(qn("w:pageBreakBefore")) is None:
                pbb = OxmlElement("w:pageBreakBefore")
                pPr.insert(0, pbb)
            # Remove the dangling empty paragraph
            el.getparent().remove(el)
            fixed += 1

        return fixed

    def save(self, output_path: str | Path) -> Path:
        if self.normalize_page_breaks_flag:
            self.normalize_page_breaks()
        out = Path(output_path).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out))
        return out


# ---------- CLI ----------

def fill_from_yaml(template: Path, content_yaml: Path, output: Path) -> None:
    with open(content_yaml, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    protections = cfg.get("protections", {}) or {}
    korean_font = protections.get("korean_font", DEFAULT_KOREAN_FONT)
    blank_between = protections.get("blank_between_paragraphs", True)
    blank_around = protections.get("blank_around_section_header", True)
    blank_around_all = protections.get("blank_around_all_section_headers", False)
    normalize_pb = protections.get("normalize_page_breaks", True)
    filler = FormFiller(template, korean_font=korean_font,
                         blank_between_paragraphs=blank_between,
                         blank_around_section_header=blank_around,
                         blank_around_all_section_headers=blank_around_all,
                         normalize_page_breaks=normalize_pb)

    # Fill table key-value pairs
    for label, value in (cfg.get("table_kv") or {}).items():
        ok = filler.fill_table_kv(str(label), str(value))
        status = "OK " if ok else "MISS"
        print(f"  [{status}] table_kv: {label!r}")

    # Replace section content (between headers)
    for header, content in (cfg.get("section_replace") or {}).items():
        ok = filler.replace_paragraphs_after(str(header), str(content))
        status = "OK " if ok else "MISS"
        print(f"  [{status}] section: {header!r}")

    # Replace single paragraph in-place (e.g., title line)
    for matcher, content in (cfg.get("paragraph_replace") or {}).items():
        ok = filler.replace_paragraph_matching(str(matcher), str(content),
                                                mode="startswith")
        status = "OK " if ok else "MISS"
        print(f"  [{status}] paragraph: {matcher!r}")

    # Document-wide pass: blank lines around ALL numbered section headers
    if blank_around_all:
        n = filler.apply_blank_around_all_section_headers()
        print(f"  [OK ] blank lines around all numbered headers: {n} inserted")

    print()
    print(filler.report())
    print()

    warnings = filler.validate()
    for w in warnings:
        print(f"  WARN: {w}")

    saved = filler.save(output)
    print(f"\nSaved: {saved}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, help="Path to template .docx")
    parser.add_argument("--content", required=True, help="Path to content YAML")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    fill_from_yaml(Path(args.template), Path(args.content), Path(args.output))


if __name__ == "__main__":
    main()
