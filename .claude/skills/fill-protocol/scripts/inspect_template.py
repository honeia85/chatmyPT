#!/usr/bin/env python3
"""Inspect a Word template — list all tables, cells, and paragraphs.

Output identifies fillable cells (likely empty after a label cell).

Usage: python3 inspect_template.py <template.docx>
"""
import sys
from pathlib import Path

from docx import Document


def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    path = Path(sys.argv[1]).expanduser().resolve()
    doc = Document(str(path))

    print(f"=== Template: {path.name} ===\n")

    print(f"Sections: {len(doc.sections)}")
    sec = doc.sections[0]
    print(
        f"  Page: {sec.page_width.cm:.1f} × {sec.page_height.cm:.1f} cm, "
        f"margins L/R/T/B: {sec.left_margin.cm:.1f}/{sec.right_margin.cm:.1f}/"
        f"{sec.top_margin.cm:.1f}/{sec.bottom_margin.cm:.1f}"
    )
    print()

    print(f"Tables: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        n_rows = len(table.rows)
        n_cols = len(table.columns)
        print(f"\n[Table {ti}] rows={n_rows}, cols={n_cols}")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell_text(cell)
                preview = text.replace("\n", " ⏎ ")
                if len(preview) > 70:
                    preview = preview[:67] + "..."
                marker = " [empty]" if not text else ""
                print(f"  ({ri},{ci}): {preview!r}{marker}")

    print(f"\nParagraphs (top-level, not in tables): {len(doc.paragraphs)}")
    for pi, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        preview = text[:80] + ("..." if len(text) > 80 else "")
        print(f"  P{pi}: {preview!r}")


if __name__ == "__main__":
    main()
